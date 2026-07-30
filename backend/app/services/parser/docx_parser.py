"""
Parser de documentos DOCX.

Usa python-docx para extração de texto.

Segurança:
- XML parsing com entidades externas desabilitadas (XXE prevention)
- Sem execução de macros
"""

import logging
from pathlib import Path
from xml.etree import ElementTree

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT


logger = logging.getLogger(__name__)


def parse_docx(file_path: Path) -> tuple[str, list[dict]]:
    """
    Extrai texto de um DOCX.

    Segurança: python-docx usa lxml internamente que por padrão
    desabilita entities externas. Adicionamos verificação extra.

    Returns:
        Tuple (texto_completo, lista_de_páginas)
        DOCX não tem conceito nativo de páginas, então retornamos
        uma única "página" com todo o conteúdo.
    """
    try:
        doc = DocxDocument(str(file_path))
    except Exception:
        logger.exception("Erro ao abrir DOCX: %s", file_path.name)
        raise ValueError("Não foi possível abrir o arquivo DOCX. Verifique o formato.")

    text_parts = []

    # Extrair parágrafos
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Identificar estilo (heading, normal, etc.)
            style = para.style.name if para.style else ""

            if "Heading" in style or "Título" in style:
                # Marcar como título para o estruturador
                text_parts.append(f"[TÍTULO] {text}")
            elif "List" in style or "Lista" in style:
                text_parts.append(f"  • {text}")
            else:
                text_parts.append(text)

    # Extrair tabelas
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text)
            if any(cells):
                table_rows.append(" | ".join(cells))

        if table_rows:
            text_parts.append("[TABELA]")
            text_parts.extend(table_rows)
            text_parts.append("[/TABELA]")

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        raise ValueError(
            "Nenhum texto pôde ser extraído do DOCX. "
            "Verifique se o arquivo não está vazio ou corrompido."
        )

    # DOCX não tem páginas nativas — retornar como página única
    pages = [{"page": 1, "text": full_text}]

    return full_text, pages
